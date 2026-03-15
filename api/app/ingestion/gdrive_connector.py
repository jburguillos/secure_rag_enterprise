"""Google Drive connector using LlamaHub (primary) with metadata/ACL enrichment."""

from __future__ import annotations

import inspect
import io
import json
import hashlib
import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from llama_index.core import Document
from llama_index.readers.google import GoogleDriveReader
from pypdf import PdfReader

from app.config import get_settings
from app.ingestion.drive_auth import oauth_credentials, service_account_credentials
from app.ingestion.path_metadata import normalize_path, path_ancestors
from app.ingestion.tabular import parse_xlsx_bytes

logger = logging.getLogger(__name__)

SUPPORTED_MIME_TYPES = {
    "application/vnd.google-apps.document",
    "application/vnd.google-apps.spreadsheet",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/plain",
}
FOLDER_MIME_TYPE = "application/vnd.google-apps.folder"


@dataclass
class DriveFile:
    file_id: str
    name: str
    mime_type: str
    web_view_link: str | None
    modified_time: str | None
    parent_folder_id: str | None = None
    folder_path: str = ""
    drive_path: str = ""


def _parse_modified_time(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None


def build_drive_service(auth_mode: str):
    settings = get_settings()
    if auth_mode == "service_account":
        if not settings.google_service_account_json:
            raise ValueError("GOOGLE_SERVICE_ACCOUNT_JSON is required for service_account mode")
        creds = service_account_credentials(settings.google_service_account_json)
    else:
        creds = oauth_credentials(settings.google_credentials_path, settings.google_token_path)
    return build("drive", "v3", credentials=creds, cache_discovery=False)


def list_drive_files(folder_id: str, service) -> tuple[list[DriveFile], list[dict[str, Any]]]:
    supported: list[DriveFile] = []
    skipped: list[dict[str, Any]] = []
    visited_folders: set[str] = set()

    def walk_folder(current_folder_id: str, current_path: str) -> None:
        if current_folder_id in visited_folders:
            return
        visited_folders.add(current_folder_id)

        query = f"'{current_folder_id}' in parents and trashed=false"
        page_token = None

        while True:
            response = (
                service.files()
                .list(
                    q=query,
                    fields="nextPageToken, files(id,name,mimeType,webViewLink,modifiedTime)",
                    includeItemsFromAllDrives=True,
                    supportsAllDrives=True,
                    corpora="allDrives",
                    pageToken=page_token,
                    pageSize=1000,
                )
                .execute()
            )

            for item in response.get("files", []):
                mime = item.get("mimeType", "")
                name = item.get("name", "")
                item_path = "/".join(part for part in [current_path, name] if part)
                item_id = item.get("id", "")

                if mime == FOLDER_MIME_TYPE and item_id:
                    walk_folder(item_id, item_path)
                    continue

                entry = DriveFile(
                    file_id=item_id,
                    name=name,
                    mime_type=mime,
                    web_view_link=item.get("webViewLink"),
                    modified_time=item.get("modifiedTime"),
                    parent_folder_id=current_folder_id,
                    folder_path=current_path,
                    drive_path=item_path or name,
                )
                if mime in SUPPORTED_MIME_TYPES and entry.file_id:
                    supported.append(entry)
                else:
                    skipped.append(
                        {
                            "file_id": entry.file_id,
                            "name": entry.name,
                            "mimeType": mime,
                            "path": normalize_path(entry.drive_path),
                        }
                    )

            page_token = response.get("nextPageToken")
            if not page_token:
                break

    walk_folder(folder_id, "")

    return supported, skipped


def fetch_permissions(file_id: str, service) -> dict[str, Any]:
    default = {
        "allowed_emails": [],
        "allowed_domains": [],
        "allowed_users": [],
        "allowed_groups": [],
        "is_public": False,
        "permissions_raw": [],
    }
    try:
        response = (
            service.permissions()
            .list(
                fileId=file_id,
                fields="permissions(id,type,emailAddress,domain,role,allowFileDiscovery)",
                supportsAllDrives=True,
            )
            .execute()
        )
    except Exception:  # noqa: BLE001
        return default

    emails: set[str] = set()
    domains: set[str] = set()
    allowed_users: set[str] = set()
    allowed_groups: set[str] = set()
    is_public = False

    permissions = response.get("permissions", [])
    for perm in permissions:
        p_type = perm.get("type")
        email = str(perm.get("emailAddress") or "").lower()
        domain = str(perm.get("domain") or "").lower()

        if p_type == "user" and email:
            emails.add(email)
            allowed_users.add(email)
        elif p_type == "group" and email:
            emails.add(email)
            allowed_groups.add(email)
        elif p_type == "domain" and domain:
            domains.add(domain)
        elif p_type == "anyone":
            is_public = True

    return {
        "allowed_emails": sorted(emails),
        "allowed_domains": sorted(domains),
        "allowed_users": sorted(allowed_users),
        "allowed_groups": sorted(allowed_groups),
        "is_public": is_public,
        "permissions_raw": permissions,
    }


def _download_request(file_id: str, mime_type: str, service):
    if mime_type == "application/vnd.google-apps.document":
        return service.files().export_media(fileId=file_id, mimeType="text/plain")
    if mime_type == "application/vnd.google-apps.spreadsheet":
        return service.files().export_media(
            fileId=file_id,
            mimeType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    return service.files().get_media(fileId=file_id, supportsAllDrives=True)


def download_drive_file_bytes(*, file_id: str, mime_type: str, service) -> bytes:
    request = _download_request(file_id=file_id, mime_type=mime_type, service=service)
    stream = io.BytesIO()
    downloader = MediaIoBaseDownload(stream, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return stream.getvalue()


def _extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    chunks: list[str] = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return "\n\n".join(chunks).strip()


def _extract_pdf_page_map(content: bytes) -> tuple[list[dict[str, Any]], int]:
    reader = PdfReader(io.BytesIO(content))
    page_map: list[dict[str, Any]] = []
    page_count = len(reader.pages)
    for page_idx, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            page_map.append({"page": page_idx, "text": text})
    return page_map, page_count


def _extract_docx_text(content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs if p.text).strip()


def _document_from_payload(*, payload: bytes, file_name: str, mime_type: str) -> tuple[str, dict[str, Any]]:
    binary_hash = hashlib.sha256(payload).hexdigest()
    if mime_type == "application/pdf":
        return _extract_pdf_text(payload), {"hash": binary_hash}
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx_text(payload), {"hash": binary_hash}
    if mime_type in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.google-apps.spreadsheet",
    }:
        parsed = parse_xlsx_bytes(payload, workbook_name=file_name)
        return parsed.document_text, {
            "hash": binary_hash,
            "sheet_map": parsed.sheet_map,
            "tabular_nodes": parsed.tabular_nodes,
            "tabular_warnings": parsed.warnings,
            "tabular_truncated": parsed.truncated,
        }
    return payload.decode("utf-8", errors="ignore"), {"hash": binary_hash}


def _download_file_content(*, file_id: str, mime_type: str, service, file_name: str) -> tuple[str, dict[str, Any]]:
    payload = download_drive_file_bytes(file_id=file_id, mime_type=mime_type, service=service)
    return _document_from_payload(payload=payload, file_name=file_name, mime_type=mime_type)


def _reader_instance(auth_mode: str):
    settings = get_settings()
    options: list[dict[str, Any]] = []
    if auth_mode == "service_account" and settings.google_service_account_json:
        options.extend(
            [
                {"service_account_key_file": settings.google_service_account_json},
                {"service_account_key": json.loads(Path(settings.google_service_account_json).read_text(encoding="utf-8"))},
            ]
        )
    else:
        options.extend(
            [
                {
                    "credentials_path": settings.google_credentials_path,
                    "token_path": settings.google_token_path,
                },
                {
                    "client_config": json.loads(Path(settings.google_credentials_path).read_text(encoding="utf-8")),
                    "authorized_user_info": json.loads(Path(settings.google_token_path).read_text(encoding="utf-8")),
                },
            ]
        )
    options.append({})

    errors: list[str] = []
    for kwargs in options:
        try:
            sig = inspect.signature(GoogleDriveReader)
            filtered = {k: v for k, v in kwargs.items() if k in sig.parameters}
            return GoogleDriveReader(**filtered)
        except Exception as exc:  # noqa: BLE001
            errors.append(str(exc))
    raise RuntimeError("Unable to initialize GoogleDriveReader: " + " | ".join(errors))


def _reader_load(reader: GoogleDriveReader, *, folder_id: str, file_ids: list[str], auth_mode: str) -> list[Document]:
    settings = get_settings()
    kwargs_options: list[dict[str, Any]] = [
        {"folder_id": folder_id, "file_ids": file_ids, "supportsAllDrives": True, "includeItemsFromAllDrives": True},
        {"folder_id": folder_id, "file_ids": file_ids},
    ]
    if auth_mode == "service_account" and settings.google_service_account_json:
        kwargs_options.insert(
            0,
            {
                "folder_id": folder_id,
                "file_ids": file_ids,
                "service_account_key_file": settings.google_service_account_json,
                "supportsAllDrives": True,
                "includeItemsFromAllDrives": True,
            },
        )

    errors: list[str] = []
    for kwargs in kwargs_options:
        try:
            sig = inspect.signature(reader.load_data)
            filtered = {k: v for k, v in kwargs.items() if k in sig.parameters}
            docs = reader.load_data(**filtered)
            return list(docs)
        except Exception as exc:  # noqa: BLE001
            errors.append(str(exc))

    logger.warning("GoogleDriveReader failed for folder=%s: %s", folder_id, " | ".join(errors))
    return []


def load_drive_documents(folder_id: str, auth_mode: str, service: Any | None = None) -> tuple[list[Document], list[dict[str, Any]]]:
    """Load drive docs and enrich with ACL metadata.

    Returns documents and skipped-file metadata.
    """

    settings = get_settings()
    drive_service = service or build_drive_service(auth_mode)
    supported_files, skipped = list_drive_files(folder_id, drive_service)
    if not supported_files:
        return [], skipped

    reader_docs: list[Document] = []
    if settings.google_drive_use_reader:
        file_ids = [item.file_id for item in supported_files]
        try:
            reader = _reader_instance(auth_mode)
            reader_docs = _reader_load(reader, folder_id=folder_id, file_ids=file_ids, auth_mode=auth_mode)
        except Exception as exc:  # noqa: BLE001
            logger.warning("GoogleDriveReader init failed for folder=%s: %s", folder_id, exc)
            reader_docs = []
    else:
        logger.info("Using native Google Drive downloader for folder=%s (%d supported files)", folder_id, len(supported_files))

    doc_by_file_id: dict[str, Document] = {}
    for doc in reader_docs:
        md = dict(doc.metadata or {})
        file_id = md.get("file_id") or md.get("id") or md.get("source_id")
        if file_id:
            doc_by_file_id[str(file_id)] = doc

    output: list[Document] = []

    for file_meta in supported_files:
        doc = doc_by_file_id.get(file_meta.file_id)
        extra_metadata: dict[str, Any] = {}
        if doc is None:
            try:
                content, extra_metadata = _download_file_content(
                    file_id=file_meta.file_id,
                    mime_type=file_meta.mime_type,
                    service=drive_service,
                    file_name=file_meta.name,
                )
                doc = Document(text=content, metadata=extra_metadata)
            except Exception as exc:  # noqa: BLE001
                skipped.append(
                    {
                        "file_id": file_meta.file_id,
                        "name": file_meta.name,
                        "mimeType": file_meta.mime_type,
                        "reason": f"content_download_failed: {exc}",
                    }
                )
                continue
        elif file_meta.mime_type in {
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.google-apps.spreadsheet",
        } and not (doc.metadata or {}).get("tabular_nodes"):
            try:
                content, extra_metadata = _download_file_content(
                    file_id=file_meta.file_id,
                    mime_type=file_meta.mime_type,
                    service=drive_service,
                    file_name=file_meta.name,
                )
                doc.text = content
                doc.metadata = {**dict(doc.metadata or {}), **extra_metadata}
            except Exception as exc:  # noqa: BLE001
                skipped.append(
                    {
                        "file_id": file_meta.file_id,
                        "name": file_meta.name,
                        "mimeType": file_meta.mime_type,
                        "reason": f"spreadsheet_parse_failed: {exc}",
                    }
                )
                continue

        if file_meta.mime_type == "application/pdf":
            try:
                pdf_bytes = download_drive_file_bytes(
                    file_id=file_meta.file_id,
                    mime_type=file_meta.mime_type,
                    service=drive_service,
                )
                page_map, page_count = _extract_pdf_page_map(pdf_bytes)
                extra_metadata["page_map"] = page_map
                extra_metadata["page_count"] = page_count
                if doc is None or not (doc.text or doc.get_content()).strip():
                    doc = Document(text="\n\n".join(item["text"] for item in page_map))
            except Exception as exc:  # noqa: BLE001
                logger.warning("Unable to enrich PDF page metadata for %s: %s", file_meta.file_id, exc)

        metadata = dict(doc.metadata or {})
        permissions = fetch_permissions(file_meta.file_id, drive_service)
        folder_path = normalize_path(file_meta.folder_path)
        drive_path = normalize_path(file_meta.drive_path or file_meta.name)
        metadata.update(
            {
                "doc_id": file_meta.file_id,
                "file_id": file_meta.file_id,
                "name": file_meta.name,
                "title": metadata.get("title") or file_meta.name,
                "mimeType": file_meta.mime_type,
                "webViewLink": file_meta.web_view_link,
                "modifiedTime": file_meta.modified_time,
                "dataset_source": "google_drive",
                "root_folder_id": folder_id,
                "parent_folder_id": file_meta.parent_folder_id,
                "folder_path": folder_path,
                "drive_path": drive_path,
                "folder_ancestors": path_ancestors(folder_path),
                "path_ancestors": path_ancestors(drive_path),
                **permissions,
                **extra_metadata,
            }
        )
        doc.metadata = metadata
        doc.id_ = file_meta.file_id
        output.append(doc)

    return output, skipped


def modified_time_from_document(doc: Document) -> datetime | None:
    return _parse_modified_time((doc.metadata or {}).get("modifiedTime"))
