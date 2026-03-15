"""Local folder ingestion fallback with ACL sidecar mapping."""

from __future__ import annotations

import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import yaml
from docx import Document as DocxDocument
from llama_index.core import Document
from pypdf import PdfReader

from app.ingestion.path_metadata import normalize_path, path_ancestors
from app.ingestion.tabular import parse_xlsx_path

SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx", ".xlsx"}
_LOCAL_MIME_TYPES = {
    ".txt": ".txt",
    ".md": ".md",
    ".pdf": ".pdf",
    ".docx": ".docx",
    ".xlsx": ".xlsx",
}


def _read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    if suffix == ".xlsx":
        return parse_xlsx_path(path).document_text
    raise ValueError(f"Unsupported local file type: {path.suffix}")


def _read_pdf_page_map(path: Path) -> list[dict[str, Any]]:
    reader = PdfReader(str(path))
    page_map: list[dict[str, Any]] = []
    for page_idx, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            page_map.append({"page": page_idx, "text": text})
    return page_map


def _load_acl_sidecar(acl_sidecar_path: str) -> dict[str, Any]:
    path = Path(acl_sidecar_path)
    if not path.exists():
        return {}
    data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    if isinstance(data, dict) and "documents" in data and isinstance(data["documents"], dict):
        return data["documents"]
    if isinstance(data, dict):
        return data
    return {}


def _default_acl() -> dict[str, Any]:
    return {
        "allowed_emails": [],
        "allowed_domains": [],
        "allowed_users": [],
        "allowed_groups": [],
        "is_public": False,
        "permissions_raw": {},
    }


def _content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def load_local_documents(path: str, acl_sidecar_path: str) -> tuple[list[Document], list[dict[str, Any]]]:
    root = Path(path)
    acl_map = _load_acl_sidecar(acl_sidecar_path)
    docs: list[Document] = []
    skipped: list[dict[str, Any]] = []

    for file_path in root.rglob("*"):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_SUFFIXES:
            skipped.append({"path": str(file_path), "reason": "unsupported_suffix"})
            continue

        try:
            tabular_result = parse_xlsx_path(file_path) if file_path.suffix.lower() == ".xlsx" else None
            text = tabular_result.document_text if tabular_result else _read_text(file_path)
        except Exception as exc:  # noqa: BLE001
            skipped.append({"path": str(file_path), "reason": f"parse_failed: {exc}"})
            continue

        rel = str(file_path.relative_to(root)).replace("\\", "/")
        drive_path = normalize_path(rel)
        folder_path = normalize_path(str(Path(rel).parent).replace("\\", "/"))
        if folder_path == ".":
            folder_path = ""
        acl = dict(_default_acl())
        acl_override = acl_map.get(rel) or acl_map.get(file_path.name) or {}
        if isinstance(acl_override, dict):
            acl.update(acl_override)

        content_hash = _content_hash(text)
        if tabular_result is not None:
            content_hash = _content_hash(
                json.dumps(
                    {
                        "document_text": tabular_result.document_text,
                        "sheet_map": tabular_result.sheet_map,
                        "tabular_nodes": tabular_result.tabular_nodes,
                    },
                    sort_keys=True,
                )
            )

        metadata = {
            "doc_id": rel,
            "file_id": rel,
            "name": file_path.name,
            "title": file_path.stem,
            "mimeType": _LOCAL_MIME_TYPES.get(file_path.suffix.lower(), file_path.suffix.lower()),
            "webViewLink": None,
            "modifiedTime": datetime.fromtimestamp(file_path.stat().st_mtime, tz=timezone.utc).isoformat(),
            "dataset_source": "local_folder",
            "source_path": str(file_path),
            "folder_path": folder_path,
            "drive_path": drive_path,
            "folder_ancestors": path_ancestors(folder_path),
            "path_ancestors": path_ancestors(drive_path),
            "hash": content_hash,
            **acl,
        }

        if file_path.suffix.lower() == ".pdf":
            metadata["page_map"] = _read_pdf_page_map(file_path)
            metadata["page_count"] = len(metadata["page_map"])
        elif tabular_result is not None:
            metadata["sheet_map"] = tabular_result.sheet_map
            metadata["tabular_nodes"] = tabular_result.tabular_nodes
            metadata["tabular_warnings"] = tabular_result.warnings
            metadata["tabular_truncated"] = tabular_result.truncated

        doc = Document(text=text, metadata=metadata, id_=rel)
        docs.append(doc)

    return docs, skipped
